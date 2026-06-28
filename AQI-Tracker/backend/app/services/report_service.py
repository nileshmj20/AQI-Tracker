from __future__ import annotations

import asyncio
import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Iterable

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.models.schemas import AqiResponse, ForecastResponse, StructuredInsights, WeatherResponse
from app.services.aqi_service import get_current_aqi, get_forecast
from app.services.insights_service import generate_insights
from app.services.weather_service import get_weather

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MEDIA_TYPE = "application/pdf"


async def build_docx_report(lat: float, lon: float, name: str | None = None) -> tuple[bytes, str]:
    aqi, weather, forecast = await asyncio.gather(get_current_aqi(lat, lon), get_weather(lat, lon), get_forecast(lat, lon, 7))
    insights = await generate_insights(aqi, weather, forecast)
    title = clean_title(name) or f"{lat:.4f}, {lon:.4f}"
    document = Document()
    setup_document(document)
    add_cover(document, title, lat, lon, aqi, weather)
    add_status_table(document, aqi, weather, insights)
    add_insights(document, insights)
    add_pollutants(document, aqi)
    add_charts(document, aqi, weather, forecast)
    add_footer_note(document)
    buffer = BytesIO()
    document.save(buffer)
    filename = f"AQI_Report_{slug(title)}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return buffer.getvalue(), filename

async def build_pdf_report(lat: float, lon: float, name: str | None = None) -> tuple[bytes, str]:
    aqi, weather, forecast = await asyncio.gather(get_current_aqi(lat, lon), get_weather(lat, lon), get_forecast(lat, lon, 7))
    insights = await generate_insights(aqi, weather, forecast)
    title = clean_title(name) or f"{lat:.4f}, {lon:.4f}"
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=.55 * inch, leftMargin=.55 * inch, topMargin=.55 * inch, bottomMargin=.55 * inch)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("AqiTitle", parent=styles["Heading1"], fontSize=22, leading=27, textColor=colors.HexColor("#ff5c20"), spaceAfter=12)
    h2 = ParagraphStyle("AqiHeading", parent=styles["Heading2"], fontSize=14, leading=18, textColor=colors.HexColor("#0b1b33"), spaceBefore=14, spaceAfter=8)
    body = ParagraphStyle("AqiBody", parent=styles["BodyText"], fontSize=9.4, leading=13, textColor=colors.HexColor("#172033"), spaceAfter=7)
    small = ParagraphStyle("AqiSmall", parent=styles["BodyText"], fontSize=8.2, leading=11, textColor=colors.HexColor("#5c6675"))
    current = weather.current
    rows = [
        ["US AQI", value(aqi.aqi), aqi.band.label if aqi.band else "Unknown"],
        ["Dominant pollutant", aqi.dominant_pollutant.label if aqi.dominant_pollutant else "Unavailable", ratio_text(aqi)],
        ["Temperature", f"{round(current.temperature_2m)}°C" if current and current.temperature_2m is not None else "Unavailable", current.weather_label if current else "Unknown"],
        ["Wind / humidity", f"{round(current.wind_speed_10m)} km/h" if current and current.wind_speed_10m is not None else "Unavailable", f"{round(current.relative_humidity_2m)}%" if current and current.relative_humidity_2m is not None else "Unavailable"],
    ]
    story = [
        Paragraph("AQI Tracker AI Report", h1),
        Paragraph(f"{title} · {lat:.4f}, {lon:.4f}", body),
        Paragraph(f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}", small),
        Spacer(1, 10),
        _pdf_table(rows, [1.55 * inch, 1.7 * inch, 3.0 * inch]),
        Paragraph("Smart recommendations", h2),
        Paragraph(f"<b>Situation summary:</b> {insights.summary}", body),
        Paragraph(f"<b>Health guidance:</b> {insights.health_guidance}", body),
        Paragraph(f"<b>Activity suggestion:</b> {insights.activity_suggestion}", body),
        Paragraph(f"<b>Trend explanation:</b> {insights.trend_explanation}", body),
        Paragraph("Pollutant breakdown", h2),
        _pdf_table([[p.label, value(p.value), p.unit or "", value(p.safety_ceiling), f"{round(p.ratio_percent)}%" if p.ratio_percent is not None else "Unavailable"] for p in aqi.pollutants] or [["Unavailable", "", "", "", ""]], [1.25 * inch, 1.0 * inch, .8 * inch, 1.2 * inch, 1.1 * inch], headers=["Pollutant", "Value", "Unit", "Ceiling", "%"]),
        Paragraph("Seven-day outlook", h2),
        _pdf_table(_forecast_rows(forecast, weather), [1.0 * inch, 1.0 * inch, 1.2 * inch, 2.8 * inch], headers=["Day", "Avg AQI", "Max temp", "Reading"]),
        Spacer(1, 8),
        Paragraph("Generated locally by AQI Tracker. For emergencies, follow official local advisories and health guidance.", small),
    ]
    doc.build(story)
    filename = f"AQI_AI_Report_{slug(title)}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    return buffer.getvalue(), filename


def _forecast_rows(forecast: ForecastResponse, weather: WeatherResponse) -> list[list[str]]:
    groups: dict[str, list[float]] = {}
    for point in forecast.hourly[:24 * 7]:
        key = point.time.date().isoformat()
        groups.setdefault(key, [])
        if point.aqi is not None:
            groups[key].append(float(point.aqi))
    rows = []
    for index, day in enumerate(weather.daily[:7]):
        values = groups.get(day.date, [])
        avg = round(sum(values) / len(values)) if values else "Unavailable"
        label = day.date[5:] if len(day.date) >= 10 else day.date
        temp = f"{round(day.temperature_2m_max)}°C" if day.temperature_2m_max is not None else "Unavailable"
        rows.append([label, str(avg), temp, day.weather_label or "Unknown"])
    if rows:
        return rows
    for key, values in list(groups.items())[:7]:
        avg = round(sum(values) / len(values)) if values else "Unavailable"
        rows.append([key[5:], str(avg), "Unavailable", "AQI forecast"])
    return rows or [["Unavailable", "", "", ""]]


def _pdf_table(rows: list[list[str]], col_widths: list[float], headers: list[str] | None = None) -> Table:
    data = [headers] + rows if headers else rows
    table = Table(data, colWidths=col_widths, hAlign="LEFT")
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef6ff") if headers else colors.HexColor("#f6f9fc")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#142033")),
        ("GRID", (0, 0), (-1, -1), .45, colors.HexColor("#cdd8e5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.4),
        ("LEADING", (0, 0), (-1, -1), 10.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
    ]
    table.setStyle(TableStyle(style))
    return table


def setup_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(22)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 2"].font.bold = True


def add_cover(document: Document, title: str, lat: float, lon: float, aqi: AqiResponse, weather: WeatherResponse) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("AQI Tracker Executive Report")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 92, 32)
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"{title}  |  {lat:.4f}, {lon:.4f}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(90, 100, 115)
    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.add_run(f"Generated {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}")
    document.add_paragraph()
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    cells = table.rows[0].cells
    cells[0].text = f"US AQI\n{value(aqi.aqi)}\n{aqi.band.label if aqi.band else 'Unknown'}"
    cells[1].text = f"Dominant pollutant\n{aqi.dominant_pollutant.label if aqi.dominant_pollutant else 'Unavailable'}\n{ratio_text(aqi)}"
    current = weather.current
    cells[2].text = f"Temperature\n{round(current.temperature_2m) if current and current.temperature_2m is not None else '—'}°C\n{current.weather_label if current else 'Unknown'}"
    cells[3].text = f"Wind / Humidity\n{round(current.wind_speed_10m) if current and current.wind_speed_10m is not None else '—'} km/h\nHumidity {round(current.relative_humidity_2m) if current and current.relative_humidity_2m is not None else '—'}%"
    shade_header_row(table)


def add_status_table(document: Document, aqi: AqiResponse, weather: WeatherResponse, insights: StructuredInsights) -> None:
    document.add_heading("Current status", level=2)
    current = weather.current
    rows = [
        ("AQI source", aqi.source),
        ("AQI band", aqi.band.label if aqi.band else "Unknown"),
        ("Health advice", aqi.band.advice if aqi.band else "No official band advice available."),
        ("Weather", current.weather_label if current else "Unknown"),
        ("Pressure", f"{round(current.pressure_msl)} hPa" if current and current.pressure_msl is not None else "Unavailable"),
        ("Insight mode", insights.source),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Details"
    for metric, detail in rows:
        cells = table.add_row().cells
        cells[0].text = metric
        cells[1].text = str(detail)
    shade_header_row(table)


def add_insights(document: Document, insights: StructuredInsights) -> None:
    document.add_heading("Insights and guidance", level=2)
    items = [
        ("Situation summary", insights.summary),
        ("Health guidance", insights.health_guidance),
        ("Activity suggestion", insights.activity_suggestion),
        ("Trend explanation", insights.trend_explanation),
    ]
    for title, text in items:
        p = document.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 92, 32)
        document.add_paragraph(text)


def add_pollutants(document: Document, aqi: AqiResponse) -> None:
    document.add_heading("Pollutant breakdown", level=2)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Pollutant", "Value", "Unit", "Reference ceiling", "% of ceiling"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for item in aqi.pollutants:
        cells = table.add_row().cells
        cells[0].text = item.label
        cells[1].text = value(item.value)
        cells[2].text = item.unit or ""
        cells[3].text = value(item.safety_ceiling)
        cells[4].text = f"{round(item.ratio_percent)}%" if item.ratio_percent is not None else "Unavailable"
    shade_header_row(table)


def add_charts(document: Document, aqi: AqiResponse, weather: WeatherResponse, forecast: ForecastResponse) -> None:
    document.add_heading("Charts", level=2)
    for title, image in [
        ("24-hour AQI trend", chart_aqi_trend(forecast)),
        ("Pollutant ceiling comparison", chart_pollutants(aqi)),
        ("7-day temperature forecast", chart_temperature(weather)),
    ]:
        document.add_paragraph(title).runs[0].bold = True
        document.add_picture(image, width=Inches(6.9))


def chart_aqi_trend(forecast: ForecastResponse) -> BytesIO:
    points = forecast.hourly[:24]
    labels = [p.time.strftime("%H:%M") for p in points]
    values = [p.aqi if p.aqi is not None else 0 for p in points]
    fig, ax = plt.subplots(figsize=(8.8, 3.1), dpi=160)
    ax.plot(labels, values, linewidth=2.4, marker="o", markersize=3)
    ax.fill_between(range(len(values)), values, alpha=0.16)
    style_axis(ax, "US AQI")
    ax.set_title("24-hour AQI trend", loc="left", fontweight="bold")
    ax.tick_params(axis="x", rotation=45, labelsize=7)
    return fig_to_bytes(fig)


def chart_pollutants(aqi: AqiResponse) -> BytesIO:
    items = [p for p in aqi.pollutants if p.ratio_percent is not None]
    labels = [p.label for p in items]
    values = [p.ratio_percent or 0 for p in items]
    fig, ax = plt.subplots(figsize=(8.8, 3.4), dpi=160)
    ax.barh(labels, values)
    style_axis(ax, "% of ceiling")
    ax.set_title("Pollutant comparison", loc="left", fontweight="bold")
    ax.invert_yaxis()
    return fig_to_bytes(fig)


def chart_temperature(weather: WeatherResponse) -> BytesIO:
    days = weather.daily
    labels = [d.date for d in days]
    highs = [d.temperature_2m_max if d.temperature_2m_max is not None else 0 for d in days]
    lows = [d.temperature_2m_min if d.temperature_2m_min is not None else 0 for d in days]
    fig, ax = plt.subplots(figsize=(8.8, 3.1), dpi=160)
    ax.plot(labels, highs, linewidth=2.4, marker="o", label="High °C")
    ax.plot(labels, lows, linewidth=2.0, marker="o", label="Low °C")
    style_axis(ax, "Temperature °C")
    ax.set_title("7-day temperature forecast", loc="left", fontweight="bold")
    ax.legend(frameon=False, loc="upper left")
    return fig_to_bytes(fig)


def style_axis(ax, ylabel: str) -> None:
    ax.set_facecolor("#f7f9fc")
    ax.grid(True, alpha=0.25)
    ax.set_ylabel(ylabel)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_alpha(0.25)
    ax.spines["bottom"].set_alpha(0.25)


def fig_to_bytes(fig) -> BytesIO:
    buffer = BytesIO()
    fig.tight_layout()
    fig.savefig(buffer, format="png", bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer


def shade_header_row(table) -> None:
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 92, 32)


def add_footer_note(document: Document) -> None:
    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Generated locally by AQI Tracker. Use official local advisories for emergency decisions.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 110, 125)


def clean_title(value: str | None) -> str | None:
    if not value:
        return None
    text = " ".join(value.strip().split())
    return text[:120] or None


def slug(value: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
    return (safe or "Selected_Location")[:48]


def value(item) -> str:
    if item is None:
        return "Unavailable"
    if isinstance(item, float):
        return f"{item:.1f}" if abs(item) < 100 else f"{item:.0f}"
    return str(item)


def ratio_text(aqi: AqiResponse) -> str:
    item = aqi.dominant_pollutant
    if not item or item.ratio_percent is None:
        return "Ratio unavailable"
    return f"{round(item.ratio_percent)}% of ceiling"
